"""Excel va pretenziya generatorlari testlari (SPEC 6).

Hujjat mijoz Uzumga yuboradigan dalil. Buzuq fayl yoki yetishmayotgan
shtrix kod — bu mahsulotning va'dasi bajarilmagani.
"""
from __future__ import annotations

from datetime import date
from decimal import Decimal
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from app.db.models import DiscrepancyKind
from app.docs.claim import build_claim
from app.docs.excel import build_report
from app.docs.models import ClaimContext, ReportRow


def make_row(**kwargs: object) -> ReportRow:
    base = {
        "sku": "SKU-1",
        "barcode": "2620101-03303001001017001",
        "title": "Elore Pure uniseks parfyum suvi, 100 ml",
        "period_from": date(2026, 7, 1),
        "period_to": date(2026, 7, 31),
        "expected_qty": 125,
        "actual_qty": 120,
        "diff_qty": 5,
        "unit_cost": Decimal("50000"),
        "loss_amount": Decimal("250000"),
        "kind": DiscrepancyKind.LOST_STOCK,
        "detected_at": date(2026, 8, 1),
    }
    base.update(kwargs)
    return ReportRow(**base)  # type: ignore[arg-type]


class TestExcel:
    def test_creates_readable_file(self, tmp_path: Path) -> None:
        path = build_report([make_row()], tmp_path / "hisobot.xlsx")
        assert path.exists()

        ws = load_workbook(path).active
        assert ws["A1"].value == "SKU"
        assert ws["B1"].value == "Shtrix kod"
        assert ws["A2"].value == "SKU-1"
        assert ws["B2"].value == "2620101-03303001001017001"

    def test_total_row(self, tmp_path: Path) -> None:
        rows = [
            make_row(diff_qty=5, loss_amount=Decimal("250000")),
            make_row(sku="SKU-2", diff_qty=3, loss_amount=Decimal("90000")),
        ]
        ws = load_workbook(build_report(rows, tmp_path / "r.xlsx")).active

        total_row = len(rows) + 2
        assert ws.cell(row=total_row, column=1).value == "JAMI"
        assert ws.cell(row=total_row, column=7).value == 8
        assert ws.cell(row=total_row, column=9).value == 340000

    def test_missing_barcode_is_flagged_not_hidden(self, tmp_path: Path) -> None:
        """Shtrix kodsiz qator yashirilmaydi — seller bilishi kerak."""
        ws = load_workbook(
            build_report([make_row(barcode="")], tmp_path / "r.xlsx")
        ).active

        assert ws["B2"].value == "⚠ YO'Q"
        # Ogohlantirish izohi qo'shilgan
        note = ws.cell(row=5, column=1).value
        assert note is not None and "shtrix kod yo'q" in note

    def test_empty_report_does_not_crash(self, tmp_path: Path) -> None:
        path = build_report([], tmp_path / "bosh.xlsx")
        ws = load_workbook(path).active
        assert ws.cell(row=2, column=1).value == "JAMI"

    def test_creates_parent_directory(self, tmp_path: Path) -> None:
        path = build_report([make_row()], tmp_path / "yangi" / "ichki" / "r.xlsx")
        assert path.exists()


class TestClaim:
    @pytest.fixture
    def ctx(self) -> ClaimContext:
        return ClaimContext(
            seller_name="YaTT Aliyev A.A.",
            seller_requisites="INN 123456789, Toshkent sh.",
            shop_title="Elore Parfume",
            shop_id="125841",
            period_from=date(2026, 7, 1),
            period_to=date(2026, 7, 31),
            rows=[make_row(), make_row(sku="SKU-2", loss_amount=Decimal("90000"))],
            created_on=date(2026, 8, 6),
        )

    def test_creates_readable_file(self, ctx: ClaimContext, tmp_path: Path) -> None:
        path = build_claim(ctx, tmp_path / "pretenziya.docx")
        assert path.exists()

        text = "\n".join(p.text for p in Document(path).paragraphs)
        assert "PRETENZIYA" in text
        assert "Elore Parfume" in text
        assert "125841" in text
        assert "YaTT Aliyev A.A." in text

    def test_total_in_digits_and_words(
        self, ctx: ClaimContext, tmp_path: Path
    ) -> None:
        """SPEC 6.2: summa raqam VA so'z bilan."""
        path = build_claim(ctx, tmp_path / "p.docx")
        text = "\n".join(p.text for p in Document(path).paragraphs)

        assert "340 000,00" in text  # raqam bilan
        assert "uch yuz qirq ming" in text  # so'z bilan

    def test_table_has_all_rows(self, ctx: ClaimContext, tmp_path: Path) -> None:
        path = build_claim(ctx, tmp_path / "p.docx")
        table = Document(path).tables[0]

        assert len(table.rows) == len(ctx.rows) + 1  # sarlavha + qatorlar
        assert table.rows[0].cells[2].text == "Shtrix kod"
        assert table.rows[1].cells[1].text == "SKU-1"

    def test_missing_barcode_shown_as_dash(self, tmp_path: Path) -> None:
        ctx = ClaimContext(
            seller_name="Test",
            seller_requisites="",
            shop_title="Shop",
            shop_id="1",
            period_from=date(2026, 7, 1),
            period_to=date(2026, 7, 31),
            rows=[make_row(barcode="")],
        )
        path = build_claim(ctx, tmp_path / "p.docx")
        table = Document(path).tables[0]
        assert table.rows[1].cells[2].text == "—"

    def test_totals_computed_from_rows(self, ctx: ClaimContext) -> None:
        assert ctx.total_amount == Decimal("340000")
        assert ctx.total_qty == 10
