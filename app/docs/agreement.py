"""Qo'shimcha kelishuv (Дополнительное соглашение) — python-docx.

❗ Nima uchun kerak: pretenziyaning o'zi Uzumni to'lovga majbur qilmaydi.
Qabul qilingan haqiqiy hujjatlar to'plamida (2026-08-08 da ko'rildi)
**ikkita** fayl bor:

    1. Pretenziya      — nima yo'qolgan, qancha pul (`claim.py`)
    2. Qo'shimcha kelishuv — huquqiy asos, imzolanadi (shu modul)

Kelishuv Oferta (vositachilik shartnomasi) ning 6.10, 6.12, 6.13
bandlariga tayanadi va tovar mulk huquqi Uzumga o'tishini rasmiylashtiradi.
Usiz seller pretenziya yuborsa — javob kutib qoladi.

Matn haqiqiy hujjatdan olingan. **Bandlar tarkibini o'zgartirmang** —
bu Uzum yuristlari ko'rgan shakl.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.docs.models import ClaimContext

_BLANK = "_______________________"

# Kelishuvning 13 bandi. Asl hujjatdagi ma'no va tartib saqlangan.
CLAUSES = {
    "ru": [
        "В соответствии с пунктами 6.10, 6.12 и 6.13 Оферты, в связи с "
        "недостачей имущества Стороны 2, Сторона 1 возмещает разницу между "
        "ценой реализации Товара и комиссией по себестоимости Товара.",
        "Итоговая стоимость компенсации рассчитывается на основе "
        "установленных тарифов. Сведения о тарифах находятся в Инструкции "
        "в пункте 6.8: https://seller.uzum.uz/manual",
        "Возмещение осуществляется согласно акту расхождения, либо акту "
        "дефектации, в зависимости от повреждения или недостачи Товаров.",
        "После возмещения право собственности на поврежденный Товар "
        "переходит Стороне 1.",
        "В случае нахождения Товаров, по которым выплачено возмещение в "
        "связи с недостачей, Сторона 1 становится собственником данных "
        "Товаров.",
        "После возмещения Сторона 2 обязуется не выдвигать претензии по "
        "сумме компенсации и не оспаривать права собственности на Товары.",
        "Наименование, количество, цена реализации, сумма недостачи "
        "имущества и т.п. указывается в Претензии.",
        "В случае нахождения товара при утере и в случае отсутствия "
        "дефектов при повреждении до выплаты компенсации, состав Претензии, "
        "а именно стоимость, наименование и количество товаров, подлежит "
        "изменению Стороной 1.",
        "Стоимость компенсируемого товара может меняться Стороной 1 "
        "согласно пункту Инструкции 6.8.",
        "Сторона 1 обязуется уведомить Сторону 2 об изменениях в Претензии, "
        "направив уведомление с итоговой суммой компенсации на электронную "
        "почту Стороны 2, указанную в личном кабинете.",
        "Настоящее Соглашение является неотъемлемой частью Оферты и "
        "составлено в 2-х экземплярах, по одному для каждой из Сторон.",
        "Все остальные условия Оферты, не затронутые настоящим Соглашением, "
        "остаются неизменными.",
        "Настоящее Соглашение вступает в силу с момента его подписания и "
        "является неотъемлемой частью Оферты.",
    ],
    "uz": [
        "Oferta shartnomasining 6.10, 6.12 va 6.13-bandlariga muvofiq, "
        "2-Tomon mol-mulkining yetishmovchiligi munosabati bilan 1-Tomon "
        "Tovarning sotuv narxi va komissiya o'rtasidagi farqni qoplaydi.",
        "Kompensatsiyaning yakuniy qiymati belgilangan tariflar asosida "
        "hisoblanadi. Tariflar haqidagi ma'lumot Instruksiyaning 6.8-bandida: "
        "https://seller.uzum.uz/manual",
        "Qoplash Tovarlarning shikastlanishi yoki yetishmovchiligiga qarab "
        "farqlar dalolatnomasi yoki brak dalolatnomasi asosida amalga "
        "oshiriladi.",
        "Qoplashdan so'ng shikastlangan Tovarga bo'lgan mulk huquqi "
        "1-Tomonga o'tadi.",
        "Yetishmovchilik bo'yicha qoplash to'langan Tovarlar topilgan "
        "taqdirda, 1-Tomon ushbu Tovarlarning egasi bo'ladi.",
        "Qoplashdan so'ng 2-Tomon kompensatsiya summasi bo'yicha da'vo "
        "qo'ymaslik va Tovarlarga mulk huquqini rad etmaslik majburiyatini "
        "oladi.",
        "Tovar nomi, miqdori, sotuv narxi, yetishmovchilik summasi va "
        "shu kabilar Pretenziyada ko'rsatiladi.",
        "Yo'qolgan tovar topilgan yoki shikast aniqlanmagan taqdirda, "
        "kompensatsiya to'languniga qadar Pretenziya tarkibi — qiymati, "
        "nomi va miqdori — 1-Tomon tomonidan o'zgartirilishi mumkin.",
        "Qoplanadigan tovar qiymati 1-Tomon tomonidan Instruksiyaning "
        "6.8-bandiga muvofiq o'zgartirilishi mumkin.",
        "1-Tomon Pretenziyadagi o'zgarishlar haqida 2-Tomonni shaxsiy "
        "kabinetda ko'rsatilgan elektron pochtasiga yakuniy kompensatsiya "
        "summasi bilan xabarnoma yuborib ogohlantirish majburiyatini oladi.",
        "Ushbu Kelishuv Ofertaning ajralmas qismi bo'lib, har bir Tomon "
        "uchun bittadan, 2 nusxada tuzilgan.",
        "Ushbu Kelishuv bilan qamrab olinmagan Ofertaning boshqa barcha "
        "shartlari o'zgarishsiz qoladi.",
        "Ushbu Kelishuv imzolangan paytdan kuchga kiradi va Ofertaning "
        "ajralmas qismi hisoblanadi.",
    ],
}

TEXTS = {
    "ru": {
        "title": "Дополнительное соглашение № _______",
        "subtitle": "к Посредническому Договору № {no} от {date} (Оферта)",
        "city": "г. Ташкент",
        "preamble": (
            "ИП ООО «UZUM MARKET», в лице {blank}, действующего на основании "
            "{blank}, с одной стороны, далее Сторона 1, и {seller}, в лице "
            "(ФИО) {name}, действующего на основании {pinfl}, с другой "
            "стороны, далее Сторона 2, заключили настоящее Дополнительное "
            "соглашение к посредническому договору № {no} от {date} "
            "(далее Оферта) о нижеследующем:"
        ),
        "side1": "Сторона 1\nИП ООО «UZUM MARKET»",
        "side2": "Сторона 2\n{seller}",
        "address": "Адрес: _______________________",
        "account": "Расчетный счет: {account}",
        "mfo": "МФО: {mfo}",
        "sign": "Подпись: ____________",
    },
    "uz": {
        "title": "№ _______ Qo'shimcha kelishuv",
        "subtitle": "№ {no} {date} sanadagi Vositachilik shartnomasiga (Oferta)",
        "city": "Toshkent sh.",
        "preamble": (
            "Bir tomondan {blank} shaxsida ish yurituvchi «UZUM MARKET» MChJ "
            "(keyingi o'rinlarda 1-Tomon) va ikkinchi tomondan {name} "
            "shaxsida ish yurituvchi {seller} (keyingi o'rinlarda 2-Tomon) "
            "№ {no} {date} sanadagi vositachilik shartnomasiga (keyingi "
            "o'rinlarda Oferta) ushbu Qo'shimcha kelishuvni quyidagilar "
            "haqida tuzdilar:"
        ),
        "side1": "1-Tomon\n«UZUM MARKET» MChJ",
        "side2": "2-Tomon\n{seller}",
        "address": "Manzil: _______________________",
        "account": "Hisob raqam: {account}",
        "mfo": "MFO: {mfo}",
        "sign": "Imzo: ____________",
    },
}


def build_agreement(ctx: ClaimContext, output_path: str | Path) -> Path:
    """Qo'shimcha kelishuv hujjatini yaratadi va yo'lini qaytaradi."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    created = ctx.created_on or date.today()
    lang = ctx.lang if ctx.lang in TEXTS else "uz"
    tr = TEXTS[lang]
    req = ctx.requisites
    seller = req.entity or req.full_name or ctx.seller_name

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for text in (tr["title"], tr["subtitle"].format(
        no=req.contract_no or "______", date=req.contract_date or "__.__.____"
    )):
        par = doc.add_paragraph(text)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.runs[0].bold = True

    doc.add_paragraph()

    head = doc.add_paragraph()
    head.add_run(tr["city"])
    head.add_run("\t" * 6 + f"{created:%d.%m.%Y}")

    doc.add_paragraph()
    doc.add_paragraph(
        tr["preamble"].format(
            blank=_BLANK,
            seller=seller,
            name=req.full_name or _BLANK,
            pinfl=req.pinfl or _BLANK,
            no=req.contract_no or "______",
            date=req.contract_date or "__.__.____",
        )
    )
    doc.add_paragraph()

    for i, clause in enumerate(CLAUSES[lang], start=1):
        doc.add_paragraph(f"{i}. {clause}")

    doc.add_paragraph()

    # --- Tomonlar rekvizitlari ---
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.rows[0].cells

    left.text = tr["side1"]
    left.add_paragraph(tr["address"])
    left.add_paragraph(tr["sign"])

    right.text = tr["side2"].format(seller=seller)
    right.add_paragraph(tr["address"])
    right.add_paragraph(tr["account"].format(account=req.bank_account or _BLANK))
    right.add_paragraph(tr["mfo"].format(mfo=req.mfo or _BLANK))
    right.add_paragraph(tr["sign"])

    doc.save(path)
    return path
