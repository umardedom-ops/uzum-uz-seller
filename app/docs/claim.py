"""Pretenziya (da'vo xati) generatori — python-docx (SPEC 6.2).

Shakl **Uzum qabul qilgan haqiqiy pretenziyadan** olingan (2026-08-08 da
tekshirildi). Tuzilishi qat'iy shu tartibda:

    ОТ (ФИО) / ИП / ПИНФЛ / РАСЧЕТНЫЙ СЧЕТ / МФО
    «Претензия»
    talab matni
    jadval: nom | shtrix kod | sabab | qoplash summasi | dona | jami
    izoh: qoplash = sotuv narxi − komissiya (Instruksiya 6.8)
    yakuniy summa
    sana, imzo

⚠️ Pretenziya yolg'iz yetarli emas: Uzum to'lovni **qo'shimcha kelishuv**
(`app/docs/agreement.py`) imzolangandan keyin qiladi. Ikkalasi birga
yuboriladi.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.docs.models import ClaimContext
from app.docs.numbers import amount_in_words, format_money

# Ustunlar qabul qilingan pretenziyadagi tartibda. O'zgartirmang —
# Uzum tomoni shu ko'rinishga o'rgangan.
TABLE_HEADERS = {
    "uz": (
        "Tovar nomi",
        "Shtrix kod",
        "Sabab\n(yo'qolish/shikast)",
        "Qoplash summasi*",
        "Miqdori (dona)",
        "Jami",
    ),
    "ru": (
        "Название товара",
        "Штрих-код",
        "Причина\n(утеря/повреждение)",
        "Размер возмещения*",
        "Кол-во (шт.)",
        "Итого\n(стоим. кол-во)",
    ),
}

TEXTS = {
    "uz": {
        "title": "Pretenziya",
        "from": "KIMDAN (F.I.Sh.):",
        "entity": "Tadbirkorlik shakli:",
        "pinfl": "PINFL:",
        "account": "HISOB RAQAM:",
        "mfo": "MFO:",
        "shop": "Do'kon:",
        "period": "Tekshiruv davri:",
        "demand": (
            "Yo'qolgan/shikastlangan tovarlar uchun pul mablag'ini "
            "qaytarishingizni so'rayman:"
        ),
        "footnote": (
            "*Qoplash summasi = Tovarning haqiqiy qiymati − Marketpleys "
            "komissiyasi.\nBatafsil: Instruksiya, 6.8-band "
            "(https://seller.uzum.uz/manual)."
        ),
        "total": "To'lanishi kerak bo'lgan yakuniy summa",
        "currency": "so'm",
        "qty_total": "Jami miqdor: {qty} dona",
        "date": "Sana:",
        "sign": "Imzo: ______________________",
        "no_barcode": "—",
    },
    "ru": {
        "title": "Претензия",
        "from": "ОТ (ФИО):",
        "entity": "ИП:",
        "pinfl": "ПИНФЛ:",
        "account": "РАСЧЕТНЫЙ СЧЕТ:",
        "mfo": "МФО:",
        "shop": "Магазин:",
        "period": "Проверяемый период:",
        "demand": (
            "Прошу вас сделать возврат денежных средств за "
            "потерянные/поврежденные товары:"
        ),
        "footnote": (
            "*Размер возмещения = Действительная стоимость товара − "
            "Комиссия маркетплейса.\nПодробнее смотрите в Инструкции, "
            "пункт 6.8 (https://seller.uzum.uz/manual)."
        ),
        "total": "Итоговая сумма к выплате",
        "currency": "сум",
        "qty_total": "Общее количество: {qty} шт.",
        "date": "Дата:",
        "sign": "Подпись: ______________________",
        "no_barcode": "—",
    },
}

#: Rekvizit bo'sh bo'lsa shu chiziqcha chiqadi — seller qo'lda to'ldiradi
_BLANK = "_______________________"


def build_claim(ctx: ClaimContext, output_path: str | Path) -> Path:
    """Pretenziya hujjatini yaratadi va yo'lini qaytaradi."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    created = ctx.created_on or date.today()
    lang = ctx.lang if ctx.lang in TEXTS else "uz"
    tr = TEXTS[lang]
    req = ctx.requisites

    doc = Document()
    _set_base_style(doc)

    # --- Sarlavha: rekvizitlar (qabul qilingan shakldagi tartib) ---
    _add_requisite(doc, tr["from"], req.full_name or ctx.seller_name)
    _add_requisite(doc, tr["entity"], req.entity)
    _add_requisite(doc, tr["pinfl"], req.pinfl)
    _add_requisite(doc, tr["account"], req.bank_account)
    _add_requisite(doc, tr["mfo"], req.mfo)

    doc.add_paragraph()

    # --- Nomi ---
    title = doc.add_paragraph(tr["title"])
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    doc.add_paragraph()

    # --- Do'kon va davr (asl shaklda yo'q, lekin dalil uchun foydali) ---
    doc.add_paragraph(f"{tr['shop']} {ctx.shop_title} (ID: {ctx.shop_id})")
    doc.add_paragraph(
        f"{tr['period']} {ctx.period_from:%d.%m.%Y} — {ctx.period_to:%d.%m.%Y}"
    )

    doc.add_paragraph()
    doc.add_paragraph(tr["demand"])
    doc.add_paragraph()

    _add_table(doc, ctx, lang)

    # --- Izoh (hisob formulasi) ---
    doc.add_paragraph()
    note = doc.add_paragraph(tr["footnote"])
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(10)

    # --- Yakuniy summa ---
    doc.add_paragraph()
    total = doc.add_paragraph()
    total.add_run(f"{tr['total']} ").bold = True
    total.add_run(f"{format_money(ctx.total_amount)} {tr['currency']}").bold = True

    # Summani so'z bilan yozish faqat o'zbekcha shaklda. `amount_in_words`
    # o'zbek tilida ishlaydi — uni rus hujjatiga qo'yish xato bo'lardi
    # ("bir million ... so'm" ruscha matn ichida). Qabul qilingan haqiqiy
    # rus pretenziyasida ham summa faqat raqam bilan yozilgan.
    if lang == "uz":
        doc.add_paragraph(f"({amount_in_words(ctx.total_amount)})").runs[0].italic = True

    doc.add_paragraph(tr["qty_total"].format(qty=ctx.total_qty))

    doc.add_paragraph()
    doc.add_paragraph()

    # --- Sana va imzo ---
    doc.add_paragraph(f"{tr['date']} {created:%d.%m.%Y}")
    doc.add_paragraph(tr["sign"])

    doc.save(path)
    return path


def _add_requisite(doc: Document, label: str, value: str) -> None:
    """Rekvizit qatori. Qiymat bo'lmasa — to'ldirish uchun chiziqcha."""
    par = doc.add_paragraph()
    par.add_run(f"{label} ").bold = True
    par.add_run(value or _BLANK)


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _add_table(doc: Document, ctx: ClaimContext, lang: str) -> None:
    headers = TABLE_HEADERS[lang]
    tr = TEXTS[lang]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0].cells
    for i, name in enumerate(headers):
        header[i].text = name
        for paragraph in header[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in ctx.rows:
        cells = table.add_row().cells
        cells[0].text = row.title
        # Shtrix kodsiz tovarga da'vo qilib bo'lmaydi — ochiq belgilaymiz
        cells[1].text = row.barcode if row.has_barcode else tr["no_barcode"]
        cells[2].text = row.reason_label(lang)
        cells[3].text = format_money(row.compensation_per_unit)
        cells[4].text = str(row.diff_qty)
        cells[5].text = format_money(row.loss_amount)
